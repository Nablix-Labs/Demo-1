"""CG-005 tests.

The exit condition is "all 6 topic documents parse without errors", so the
tests that matter run against the real documents. Around those sit unit tests
for the three Word quirks that made the first version of this parser crash or
silently lose content:

  * paragraph.style can be None          (Topic 5)
  * style names vary in case             ("Normal" vs "normal")
  * bullets are not identified by style  (4 of 6 documents style them "normal")

Those three are tested directly with stub paragraphs so they still run on a
machine that does not have the source documents.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx_parser import (                       # noqa: E402
    SECTION_A_SUBSECTIONS,
    DocxParseError,
    Para,
    ParsedTopicDocument,
    Section,
    heading_level,
    is_list_item,
    parse_all_topic_documents,
    parse_topic_document,
    style_name,
)
from sources import find_topic_documents        # noqa: E402

TOPIC_DOCS = find_topic_documents()
needs_docs = pytest.mark.skipif(not TOPIC_DOCS, reason="topic documents not available")


@pytest.fixture(scope="module")
def parsed():
    if not TOPIC_DOCS:
        pytest.skip("topic documents not available")
    return parse_all_topic_documents()


# ──────────────────────────────────────────────────────────────────────
# Word quirks, tested without needing the real files
# ──────────────────────────────────────────────────────────────────────

class _StubStyle:
    def __init__(self, name):
        self.name = name


class _StubParagraph:
    """Enough of a python-docx Paragraph for the style helpers."""

    def __init__(self, style):
        self.style = style


class _RaisingParagraph:
    @property
    def style(self):
        raise ValueError("style lookup blew up")


def test_style_name_survives_a_missing_style():
    """Topic 5 has 119 paragraphs with no style. This was the original crash."""
    assert style_name(_StubParagraph(None)) == ""


def test_style_name_survives_a_style_with_no_name():
    assert style_name(_StubParagraph(_StubStyle(None))) == ""


def test_style_name_survives_a_raising_style():
    assert style_name(_RaisingParagraph()) == ""


def test_style_name_is_lowercased_and_stripped():
    """"Normal" and "normal" both appear across the six documents."""
    assert style_name(_StubParagraph(_StubStyle("  Normal  "))) == "normal"
    assert style_name(_StubParagraph(_StubStyle("normal"))) == "normal"


def test_heading_level_reads_all_three_levels_case_insensitively():
    assert heading_level(_StubParagraph(_StubStyle("Heading 1"))) == 1
    assert heading_level(_StubParagraph(_StubStyle("heading 2"))) == 2
    assert heading_level(_StubParagraph(_StubStyle("HEADING 3"))) == 3


def test_heading_level_is_none_for_body_text():
    for style in ("Normal", "normal", "List Paragraph", "Title", None):
        stub = _StubParagraph(_StubStyle(style) if style else None)
        assert heading_level(stub) is None


# ──────────────────────────────────────────────────────────────────────
# Section and document helpers
# ──────────────────────────────────────────────────────────────────────

def test_section_list_items_ignores_plain_paragraphs():
    sec = Section(heading="Included", level=2, paragraphs=[
        Para("a lead-in line", is_list_item=False),
        Para("first item", is_list_item=True),
        Para("second item", is_list_item=True),
    ])
    assert sec.list_items == ["first item", "second item"]
    assert sec.text.startswith("a lead-in line")


def test_subsection_lookup_ignores_case_and_space():
    inner = Section(heading="Learning Goal", level=2)
    outer = Section(heading="A. Internal Concept Sheet", level=1,
                    subsections={"Learning Goal": inner})
    assert outer.subsection("learning goal") is inner
    assert outer.subsection("  LEARNING GOAL ") is inner
    assert outer.subsection("Mindset Shift") is None


def test_topic_title_strips_the_prefix():
    doc = ParsedTopicDocument(Path("Topic_1_Formatted.docx"),
                              title="Topic 1 — What Is Algebra?", topic_id="ALG-ORI-01")
    assert doc.topic_title == "What Is Algebra?"
    assert doc.topic_number == 1


def test_topic_title_accepts_a_plain_hyphen():
    doc = ParsedTopicDocument(Path("x.docx"), title="Topic 4 - Expressions", topic_id=None)
    assert doc.topic_title == "Expressions"
    assert doc.topic_number == 4


def test_topic_number_falls_back_to_the_topic_id():
    doc = ParsedTopicDocument(Path("x.docx"), title="Untitled", topic_id="ALG-ORI-06")
    assert doc.topic_number == 6


# ──────────────────────────────────────────────────────────────────────
# Errors
# ──────────────────────────────────────────────────────────────────────

def test_missing_file_raises_a_clear_error(tmp_path):
    with pytest.raises(DocxParseError, match="No such topic document"):
        parse_topic_document(tmp_path / "nope.docx")


def test_unreadable_file_raises_a_clear_error(tmp_path):
    bad = tmp_path / "not-really.docx"
    bad.write_text("this is not a zip archive")
    with pytest.raises(DocxParseError, match="Could not open"):
        parse_topic_document(bad)


def test_no_documents_found_raises_rather_than_passing_vacuously():
    """An empty list would let the exit-condition test pass reading nothing."""
    with pytest.raises(DocxParseError, match="No topic documents found"):
        parse_all_topic_documents(paths=[])


# ──────────────────────────────────────────────────────────────────────
# The exit condition: all six real documents
# ──────────────────────────────────────────────────────────────────────

@needs_docs
def test_all_six_documents_are_present():
    assert len(TOPIC_DOCS) == 6, [p.name for p in TOPIC_DOCS]


@needs_docs
def test_all_six_documents_parse_without_errors(parsed):
    """CG-005 exit condition."""
    assert len(parsed) == 6


@needs_docs
def test_every_document_yields_a_topic_id_number_and_title(parsed):
    for i, doc in enumerate(parsed, start=1):
        assert doc.topic_id == f"ALG-ORI-{i:02d}", doc.source_file_name
        assert doc.topic_number == i, doc.source_file_name
        assert doc.topic_title, doc.source_file_name
        assert not doc.topic_title.lower().startswith("topic ")


@needs_docs
def test_every_document_has_both_top_level_sections(parsed):
    for doc in parsed:
        assert doc.concept_sheet is not None, doc.source_file_name
        assert doc.designer_handoff is not None, doc.source_file_name


@needs_docs
def test_section_a_has_the_same_seven_subsections_everywhere(parsed):
    """Section A is uniform across all six; Section B is not."""
    for doc in parsed:
        for want in SECTION_A_SUBSECTIONS:
            sub = doc.concept_sheet.subsection(want)
            assert sub is not None, f"{doc.source_file_name} missing {want!r}"
            assert sub.paragraphs, f"{doc.source_file_name} has {want!r} empty"


@needs_docs
def test_section_a_subsections_are_in_document_order(parsed):
    for doc in parsed:
        found = [k.strip() for k in doc.concept_sheet.subsections]
        assert found == list(SECTION_A_SUBSECTIONS), doc.source_file_name


@needs_docs
def test_scope_subsections_are_detected_as_lists_in_every_document(parsed):
    """The point of using numPr instead of the style name.

    Topics 2, 3, 4 and 6 style these bullets "normal", identically to body
    text, so a style-based check would report zero items for four of six.
    """
    for doc in parsed:
        for want in ("Included", "Excluded", "Misconceptions to Prevent"):
            sub = doc.concept_sheet.subsection(want)
            assert len(sub.list_items) == len(sub.paragraphs), (
                f"{doc.source_file_name} {want}: "
                f"{len(sub.list_items)} of {len(sub.paragraphs)} detected"
            )
            assert len(sub.list_items) >= 4, f"{doc.source_file_name} {want}"


@needs_docs
def test_golden_rules_are_captured_for_every_document(parsed):
    """NormalizedTopicBrief.golden_rules is fed from Section B."""
    for doc in parsed:
        rules = doc.designer_handoff.subsection("Part 2 — Golden Rules")
        assert rules is not None, doc.source_file_name
        assert len(rules.list_items) >= 5, doc.source_file_name


@needs_docs
def test_no_paragraph_text_is_dropped(parsed):
    """Every non-empty paragraph lands in the preamble or in some section.

    Guards against the heading walker quietly discarding content, which a
    structure-only assertion would not catch.
    """
    import docx
    from docx_parser import heading_level as _level, iter_body, style_name as _style
    from docx.table import Table

    for doc in parsed:
        def count(sec):
            return len(sec.paragraphs) + sum(count(s) for s in sec.subsections.values())

        kept = len(doc.preamble) + sum(count(s) for s in doc.sections.values())

        document = docx.Document(str(doc.source_path))
        expected = 0
        seen_title = False
        for item in iter_body(document):
            if isinstance(item, Table) or not item.text.strip():
                continue
            if _style(item) == "title" and not seen_title:
                seen_title = True
                continue
            if _level(item) is not None:
                continue
            expected += 1

        assert kept == expected, f"{doc.source_file_name}: kept {kept} of {expected}"


@needs_docs
def test_tables_are_attached_to_sections_not_lost(parsed):
    """python-docx lists tables separately; they must still land somewhere."""
    import docx

    for doc in parsed:
        def count(sec):
            return len(sec.tables) + sum(count(s) for s in sec.subsections.values())

        kept = sum(count(s) for s in doc.sections.values())
        actual = len(docx.Document(str(doc.source_path)).tables)
        assert kept == actual, f"{doc.source_file_name}: {kept} of {actual} tables"


@needs_docs
def test_topic_5_the_unstyled_document_parses_fully(parsed):
    """Topic 5 has 119 paragraphs with no style at all. Regression guard."""
    doc = next(d for d in parsed if d.topic_number == 5)
    assert doc.topic_title == "Terms, Coefficients and Factors"
    assert len(doc.concept_sheet.subsections) == 7
    unstyled = [p for s in doc.concept_sheet.subsections.values()
                for p in s.paragraphs if p.style == ""]
    assert unstyled, "expected unstyled paragraphs to be kept, not skipped"
