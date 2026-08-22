"""CG-005: structural parser for the formatted topic documents.

This module turns a Topic_N_Formatted.docx into a plain Python object that
mirrors the document's own structure. It does not interpret the content --
mapping the parsed sections onto NormalizedTopicBrief is CG-006's job. Keeping
the two apart means the awkward Word-specific handling lives in one place.

Document shape (identical across all six topics):

    Title                 "Topic 1 - What Is Algebra?"
    <preamble>            includes "Topic ID: ALG-ORI-01"
    Heading 1             A. Internal Concept Sheet
      Heading 2             Learning Goal, Mindset Shift, Core Message,
                            Main Visual Metaphor, Included, Excluded,
                            Misconceptions to Prevent
    Heading 1             B. Designer Handoff
      Heading 2             Part 1 - Creative Brief, Part 2 - Golden Rules,
                            Part 3 - Storyboard and Script, ...
        Heading 3           Creative Goal / Scene 1 / Checkpoint 1 / ...

Section A is the same seven subsections in the same order in every document.
Section B varies: Topic 1 has 14 H2s, the rest have 16.

Three things about these files break the obvious implementation, so they are
handled deliberately:

1. `paragraph.style` can be None. Topic 5 has 119 such paragraphs. Reading
   `p.style.name` directly raises AttributeError, which is how this parser
   first failed.

2. Style names are not consistently cased. All of "Normal", "normal",
   "List Paragraph" and no style at all appear across the six documents, so
   every style comparison here is lowercased.

3. Bullets cannot be identified by style. Topics 1 and 5 style them
   "List Paragraph"; Topics 2, 3, 4 and 6 style them "normal" -- the same
   style as ordinary body text. The reliable signal is the w:numPr element in
   the paragraph's XML, which is present on every genuine list item in all six
   documents and absent everywhere else. `Para.is_list_item` uses that.

Tables are attached to the section they appear in. python-docx exposes
`document.paragraphs` and `document.tables` as separate flat lists with no
ordering between them, so the body is walked element by element instead.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator, Optional

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from sources import find_topic_documents

# Style names, lowercased, as compared against below.
TITLE_STYLE = "title"
HEADING_STYLES = {"heading 1": 1, "heading 2": 2, "heading 3": 3}

TOPIC_ID_RE = re.compile(r"topic\s*id\s*[:\-]\s*([A-Za-z0-9\-]+)", re.IGNORECASE)

SECTION_A_PREFIX = "a."
SECTION_B_PREFIX = "b."

# Section A subsections, in document order. Every topic has all seven.
SECTION_A_SUBSECTIONS = (
    "Learning Goal",
    "Mindset Shift",
    "Core Message",
    "Main Visual Metaphor",
    "Included",
    "Excluded",
    "Misconceptions to Prevent",
)


class DocxParseError(Exception):
    """The document could not be read, or is missing required structure."""


# ──────────────────────────────────────────────────────────────────────
# Word-level helpers
# ──────────────────────────────────────────────────────────────────────

def style_name(paragraph: Paragraph) -> str:
    """Lowercased style name, or "" when the paragraph has no style.

    Both fallbacks matter. See note 1 and note 2 in the module docstring.
    """
    try:
        style = paragraph.style
    except Exception:
        return ""
    if style is None or style.name is None:
        return ""
    return style.name.strip().lower()


def is_list_item(paragraph: Paragraph) -> bool:
    """True when Word is rendering this paragraph as a bullet or number.

    Checks for w:numPr (numbering properties) in the paragraph's XML rather
    than trusting the style name, for the reason in note 3 above.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def heading_level(paragraph: Paragraph) -> Optional[int]:
    """1, 2 or 3 for a heading paragraph; None for anything else."""
    return HEADING_STYLES.get(style_name(paragraph))


def iter_body(document: docx.document.Document) -> Iterator[object]:
    """Yield Paragraph and Table objects in the order they appear.

    `document.paragraphs` and `document.tables` are separate lists, so a table
    cannot be attributed to the heading above it using those alone. Walking the
    body's XML children preserves the real order.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


# ──────────────────────────────────────────────────────────────────────
# Parsed structures
# ──────────────────────────────────────────────────────────────────────

@dataclass
class Para:
    """One paragraph of body text."""

    text: str
    is_list_item: bool = False
    style: str = ""

    def __str__(self) -> str:
        return self.text


@dataclass
class Section:
    """A heading and everything under it, up to the next heading of its level.

    Used for H1, H2 and H3 alike. `paragraphs` holds text that sits directly
    under this heading; text under a deeper heading belongs to that subsection
    instead.
    """

    heading: str
    level: int
    paragraphs: list[Para] = field(default_factory=list)
    subsections: dict[str, "Section"] = field(default_factory=dict)
    tables: list[list[list[str]]] = field(default_factory=list)

    @property
    def text(self) -> str:
        """All paragraphs directly under this heading, joined by newlines."""
        return "\n".join(p.text for p in self.paragraphs)

    @property
    def list_items(self) -> list[str]:
        """Just the bulleted or numbered lines."""
        return [p.text for p in self.paragraphs if p.is_list_item]

    def subsection(self, name: str) -> Optional["Section"]:
        """Look up a subsection by name, ignoring case and surrounding space."""
        wanted = name.strip().lower()
        for key, sub in self.subsections.items():
            if key.strip().lower() == wanted:
                return sub
        return None

    def __repr__(self) -> str:
        return (
            f"Section(heading={self.heading!r}, level={self.level}, "
            f"paras={len(self.paragraphs)}, subs={len(self.subsections)}, "
            f"tables={len(self.tables)})"
        )


@dataclass
class ParsedTopicDocument:
    """Everything CG-005 pulls out of one topic document."""

    source_path: Path
    title: str
    topic_id: Optional[str]
    preamble: list[Para] = field(default_factory=list)
    sections: dict[str, Section] = field(default_factory=dict)

    @property
    def source_file_name(self) -> str:
        return self.source_path.name

    @property
    def topic_number(self) -> Optional[int]:
        """The N in "Topic N - ...", taken from the title.

        Falls back to the trailing digits of the topic id, then to the
        filename, so a retitled document still resolves.
        """
        for candidate in (self.title, self.topic_id or "", self.source_file_name):
            match = re.search(r"(?:topic[_\s]*)?(\d+)", candidate, re.IGNORECASE)
            if match:
                return int(match.group(1))
        return None

    @property
    def topic_title(self) -> str:
        """Title with the "Topic N - " prefix removed.

        The documents use an em dash; a plain hyphen is accepted too in case a
        future document is typed differently.
        """
        cleaned = re.sub(
            r"^\s*topic\s*\d+\s*[—–\-]\s*", "", self.title, flags=re.IGNORECASE
        )
        return cleaned.strip()

    def section(self, prefix: str) -> Optional[Section]:
        """Section A or B, found by its heading prefix."""
        wanted = prefix.strip().lower()
        for key, sec in self.sections.items():
            if key.strip().lower().startswith(wanted):
                return sec
        return None

    @property
    def concept_sheet(self) -> Optional[Section]:
        """Section A - Internal Concept Sheet."""
        return self.section(SECTION_A_PREFIX)

    @property
    def designer_handoff(self) -> Optional[Section]:
        """Section B - Designer Handoff."""
        return self.section(SECTION_B_PREFIX)

    def __repr__(self) -> str:
        return (
            f"ParsedTopicDocument({self.source_file_name!r}, "
            f"topic_id={self.topic_id!r}, sections={list(self.sections)})"
        )


# ──────────────────────────────────────────────────────────────────────
# Parsing
# ──────────────────────────────────────────────────────────────────────

def _table_to_rows(table: Table) -> list[list[str]]:
    """A table as a list of rows of cell text."""
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def parse_topic_document(path: str | Path) -> ParsedTopicDocument:
    """Parse one topic document into its heading tree.

    Raises DocxParseError if the file is missing or unreadable. A document that
    opens but lacks an expected heading is returned as-is; checking for
    mandatory sections is CG-007's job, so that this stays a faithful reader.
    """
    path = Path(path)
    if not path.is_file():
        raise DocxParseError(f"No such topic document: {path}")

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocxParseError(f"Could not open {path.name}: {exc}") from exc

    parsed = ParsedTopicDocument(source_path=path, title="", topic_id=None)

    # The heading currently being filled at each level. current[0] is the H1,
    # current[1] the H2 below it, current[2] the H3 below that. A new heading
    # at level N replaces current[N-1] and clears everything deeper.
    current: list[Optional[Section]] = [None, None, None]

    def deepest() -> Optional[Section]:
        for section in reversed(current):
            if section is not None:
                return section
        return None

    for item in iter_body(document):
        if isinstance(item, Table):
            target = deepest()
            rows = _table_to_rows(item)
            if target is not None:
                target.tables.append(rows)
            continue

        text = item.text.strip()
        style = style_name(item)

        if style == TITLE_STYLE:
            # First Title paragraph wins; later ones are treated as body text.
            if not parsed.title and text:
                parsed.title = text
                continue

        level = heading_level(item)
        if level is not None and text:
            section = Section(heading=text, level=level)
            if level == 1:
                parsed.sections[text] = section
            else:
                parent = current[level - 2]
                if parent is None:
                    # A heading with no parent above it. Promote it rather than
                    # dropping content, so nothing is silently lost.
                    parsed.sections[text] = section
                else:
                    parent.subsections[text] = section
            current[level - 1] = section
            for deeper in range(level, 3):
                current[deeper] = None
            continue

        if not text:
            continue

        para = Para(text=text, is_list_item=is_list_item(item), style=style)

        if parsed.topic_id is None:
            match = TOPIC_ID_RE.search(text)
            if match:
                parsed.topic_id = match.group(1).strip()

        target = deepest()
        if target is None:
            parsed.preamble.append(para)
        else:
            target.paragraphs.append(para)

    if not parsed.title:
        # Every current document has a Title paragraph, but falling back to the
        # first preamble line is better than returning an empty title.
        if parsed.preamble:
            parsed.title = parsed.preamble[0].text

    return parsed


def parse_all_topic_documents(paths: Optional[list[Path]] = None) -> list[ParsedTopicDocument]:
    """Parse every topic document found by `sources`, in filename order.

    Raises DocxParseError if no documents are found, since a silently empty
    list would let the CG-005 tests pass without reading anything.
    """
    paths = list(paths) if paths is not None else find_topic_documents()
    if not paths:
        from sources import describe_sources

        raise DocxParseError("No topic documents found.\n" + describe_sources())
    return [parse_topic_document(p) for p in paths]


if __name__ == "__main__":
    for doc in parse_all_topic_documents():
        concept = doc.concept_sheet
        handoff = doc.designer_handoff
        print(f"\n{doc.source_file_name}")
        print(f"  title      : {doc.title}")
        print(f"  topic_id   : {doc.topic_id}   number={doc.topic_number}")
        print(f"  topic_title: {doc.topic_title}")
        print(f"  preamble   : {len(doc.preamble)} paras")
        for name, sec in doc.sections.items():
            print(f"  [{name}] {len(sec.subsections)} subsections")
        if concept is not None:
            for want in SECTION_A_SUBSECTIONS:
                sub = concept.subsection(want)
                if sub is None:
                    print(f"    MISSING: {want}")
                else:
                    print(
                        f"    {want:28} {len(sub.paragraphs):2} paras, "
                        f"{len(sub.list_items):2} list items"
                    )
