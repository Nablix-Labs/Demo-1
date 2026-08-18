from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/tacticalcamel/Desktop/Nablix")
OUTPUT = ROOT / "output" / "doc" / "Saravanan_Schema_3_0_Pending_Work.docx"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_code(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5F7")
    set_cell_margins(cell, 90, 120, 90, 120)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(7.6)
    run.font.color.rgb = RGBColor(31, 41, 55)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_numbered_item(document: Document, title: str, body: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    paragraph.add_run(f" - {body}")


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(8.8)
    styles["Normal"].paragraph_format.space_after = Pt(3)
    styles["Normal"].paragraph_format.line_spacing = 1.02
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(21)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(15, 55, 88)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(12.5)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(15, 76, 92)
    styles["Heading 1"].paragraph_format.space_before = Pt(7)
    styles["Heading 1"].paragraph_format.space_after = Pt(3)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(10.3)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(24, 92, 115)
    styles["Heading 2"].paragraph_format.space_before = Pt(5)
    styles["Heading 2"].paragraph_format.space_after = Pt(2)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    title.add_run("Schema 3.0 - Saravanan's Remaining Work")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("Student Model handoff | Based on mathtutor-student master 32b721f4")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(75, 85, 99)

    status = document.add_table(rows=1, cols=1)
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = status.cell(0, 0)
    set_cell_shading(cell, "E8F3F5")
    set_cell_margins(cell, 100, 140, 100, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Goal: ")
    r.bold = True
    p.add_run(
        "prevent two devices or retries from advancing the same learning journey differently. "
        "The Tutor Backend already sends the required fields. The Student Model must validate and use them."
    )

    document.add_heading("What is missing today", level=1)
    add_bullet(document, "SessionEventIn accepts extra fields but does not define or validate source_turn_id and expected_journey_version.")
    add_bullet(document, "The router stores request_id as source_turn_id, so the real browser/canvas turn is lost.")
    add_bullet(document, "The current journey version is never compared with the caller's expected version.")
    add_bullet(document, "A duplicate request returns DUPLICATE_REQUEST instead of safely replaying the already-created result.")

    document.add_heading("Required changes", level=1)
    add_numbered_item(
        document,
        "Define the two fields",
        "In app/schemas/session.py, add them to SessionEventIn. Require them for every event except SESSION_OPENED and DIAGNOSTIC_QUESTION_SET_REQUESTED.",
    )
    add_code(
        document,
        "class SessionEventIn(BaseModel):\n"
        "    request_id: str\n"
        "    event_type: str\n"
        "    topic_id: str\n"
        "    student_id: str\n"
        "    source_turn_id: str | None = None\n"
        "    expected_journey_version: int | None = None",
    )

    add_numbered_item(
        document,
        "Check the version before changing state",
        "In post_session_event(), first lock the student's current mastery_progress row. Keep it locked through duplicate lookup, version check, dispatch(), and save in one transaction. This also queues two simultaneous copies of the same request.",
    )
    add_code(
        document,
        "if body.expected_journey_version != current_journey.get(\"version\"):\n"
        "    raise AppError(\n"
        "        \"Journey changed; refresh before retrying.\",\n"
        "        error_code=\"JOURNEY_VERSION_CONFLICT\", status_code=409,\n"
        "        details={\"current_journey_state\": current_journey},\n"
        "    )",
    )

    add_numbered_item(
        document,
        "Return the current state on conflict",
        "Extend AppError and app_error_handler() in app/api/errors.py so the 409 body includes the complete current journey_state object (the JSON below is shortened). Nothing may be written on this path.",
    )
    add_code(
        document,
        "{\n"
        "  \"error_code\": \"JOURNEY_VERSION_CONFLICT\",\n"
        "  \"message\": \"Journey changed; refresh before retrying.\",\n"
        "  \"current_journey_state\": {\"version\": 8, \"current_phase\": \"PHASE_3_INDEPENDENT_PRACTICE\"}\n"
        "}",
    )

    add_numbered_item(
        document,
        "Store the real turn ID",
        "In app/api/routers/session.py, create ProcessedEvent with source_turn_id=body.source_turn_id, not body.request_id. The existing database column already supports this; no new column is needed.",
    )
    add_code(
        document,
        "ProcessedEvent(\n"
        "    event_id=body.request_id,\n"
        "    source_turn_id=body.source_turn_id,\n"
        "    event_type=body.event_type, ...\n"
        ")\n"
        "# After dispatch succeeds:\n"
        "ledger.response_body = out.model_dump(mode=\"json\")",
    )

    add_numbered_item(
        document,
        "Make retries safe",
        "After acquiring the mastery_progress lock, check request_id before the version. Add response_body JSONB to ProcessedEvent and save the successful SessionEventOut. A queued duplicate then returns SessionEventOut.model_validate(response_body) with HTTP 200. Never dispatch twice.",
    )

    document.add_page_break()
    document.add_heading("Example request from Tutor Backend", level=1)
    add_code(
        document,
        "{\n"
        "  \"request_id\": \"SESSION123:TURN456:INCORRECT_ATTEMPT\",\n"
        "  \"event_type\": \"INCORRECT_ATTEMPT\",\n"
        "  \"topic_id\": \"ALG-ORI-02\", \"student_id\": \"ST001\",\n"
        "  \"source_turn_id\": \"TURN456\",\n"
        "  \"expected_journey_version\": 7,\n"
        "  \"question_id\": \"Q-T02-004\", \"micro_skill_ids\": [\"MS-01\"],\n"
        "  \"student_response\": \"x = 4\", \"error_code\": \"SIGN_ERROR\"\n"
        "}",
    )

    document.add_heading("Files and functions to touch", level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "File"
    headers[1].text = "Small change"
    for header in headers:
        set_cell_shading(header, "0F4C5C")
        header.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in header.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    rows = (
        ("app/schemas/session.py", "Add the two request fields and conditional validation."),
        ("app/repositories/journey_repo.py", "Add load_for_update() so version check and write cannot overlap."),
        ("app/api/errors.py", "Allow conflict details in the JSON response."),
        ("app/models/student_model.py + migration", "Add ProcessedEvent.response_body JSONB for durable retry replay."),
        ("app/api/routers/session.py", "Replay duplicate, lock/load, compare version, store source turn, dispatch, save response."),
        ("tests/", "Add four focused integration checks listed below."),
    )
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_margins(cell, 55, 75, 55, 75)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.8)

    document.add_heading("Done when these four checks pass", level=1)
    add_bullet(document, "Correct expected version: event succeeds and journey version advances once.")
    add_bullet(document, "Stale expected version: 409 JOURNEY_VERSION_CONFLICT returns the current journey; nothing changes.")
    add_bullet(document, "Same request_id twice: the second call returns the same Schema response with HTTP 200 and does not advance again.")
    add_bullet(document, "ProcessedEvent.source_turn_id equals the supplied source_turn_id.")

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Out of scope: ")
    run.bold = True
    note.add_run(
        "no explicit hint-request event is needed. Canvas, text, and voice already send the same Schema 3.0 attempt events."
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Nablix - Schema 3.0 Student Model handoff")
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(107, 114, 128)
    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
